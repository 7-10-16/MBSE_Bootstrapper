import os
import docx
import openpyxl

def generate_docx_file(file_path, content_list):
    """
    Generate a .docx file with the specified content.

    Args:
        file_path (str): The path to the .docx file to be created.
        content_list (list): A list of content elements, each element can be either a dictionary
            with keys 'type' (e.g., 'heading', 'paragraph') and 'text' (the content text),
            or a plain text string.

    Returns:
        None
    """
    doc = docx.Document()
    for content in content_list:
        if isinstance(content, str):  # Check if it's a string
            doc.add_paragraph(content)
        elif isinstance(content, dict):  # Check if it's a dictionary
            if content['type'] == 'heading':
                doc.add_heading(content['text'], level=1)
            elif content['type'] == 'paragraph':
                doc.add_paragraph(content['text'])
    doc.save(file_path)

def generate_txt_file(file_path, content_list):
    """
    Generate a .txt file with the specified content.

    Args:
        file_path (str): The path to the .txt file to be created.
        content_list (list): A list of content lines.

    Returns:
        None
    """
    with open(file_path, 'w') as txt_file:
        for line in content_list:
            txt_file.write(line + '\n')

def generate_excel_file(file_path, content_list):
    """
    Generate an .xlsx file with the specified content.

    Args:
        file_path (str): The path to the .xlsx file to be created.
        content_list (list): A list of rows, each row is represented as a list of values.

    Returns:
        None
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    for row_data in content_list:
        sheet.append(row_data)
    workbook.save(file_path)
